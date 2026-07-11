"""
Regenerate the Word (.docx) client delivery SOP from its markdown source.

Usage:  python scripts/md2docx.py
Requires: pip install python-docx

Run this any time docs/CLIENT-DELIVERY-SOP.md changes, so the two files
never drift apart. The markdown file is the source of truth — always edit
that, then regenerate the .docx from it, never the other way around.
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x0E, 0x1B, 0x2A)
GOLD = RGBColor(0xB5, 0x8A, 0x1A)
SLATE = RGBColor(0x33, 0x41, 0x55)
MUTE = RGBColor(0x5B, 0x66, 0x75)
LINE_GOLD = "C9A227"

REPO_ROOT = Path(__file__).resolve().parent.parent
SRC = REPO_ROOT / "docs" / "CLIENT-DELIVERY-SOP.md"
OUT = REPO_ROOT / "docs" / "The Trades Lab - Client Delivery SOP.docx"

def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_bottom_border(paragraph, color=LINE_GOLD, sz=10):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_inline_runs(paragraph, text, base_size=10.5, color=None, base_bold=False):
    """Parse **bold** and `code` inline markdown into runs."""
    tokens = re.split(r'(\*\*.*?\*\*|`[^`]+`)', text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            run = paragraph.add_run(tok[2:-2])
            run.bold = True
        elif tok.startswith('`') and tok.endswith('`'):
            run = paragraph.add_run(tok[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(base_size - 0.5)
            run.font.color.rgb = RGBColor(0x8A, 0x5A, 0x00)
        else:
            run = paragraph.add_run(tok)
        run.font.size = Pt(base_size)
        if color:
            run.font.color.rgb = color
        if base_bold:
            run.bold = True

def new_doc():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    style.font.color.rgb = SLATE
    section = doc.sections[0]
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    return doc

doc = new_doc()

# ── Title page ──────────────────────────────────────────────
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = title_p.add_run('THE TRADES LAB')
run.font.size = Pt(11)
run.font.color.rgb = GOLD
run.bold = True
run.font.name = 'Calibri'

h = doc.add_paragraph()
run = h.add_run('Client Delivery SOP')
run.font.size = Pt(30)
run.font.color.rgb = NAVY
run.bold = True
h.space_after = Pt(4)

sub = doc.add_paragraph()
run = sub.add_run('Milestone-by-milestone checklist for delivering the WOW Website + $100k Funnel system — repeatable, team-shareable, tick-box ready.')
run.font.size = Pt(11.5)
run.font.color.rgb = MUTE
run.italic = True
add_bottom_border(sub, sz=14)
sub.space_after = Pt(14)

legend = doc.add_paragraph()
add_inline_runs(legend, 'Legend:  🖥️ VS Code   ·   🤖 Claude   ·   🔷 GHL   ·   ☁️ Deploy   ·   📊 Analytics', base_size=9.5, color=MUTE)
legend.space_after = Pt(2)

note = doc.add_paragraph()
add_inline_runs(note, 'This product is built to be reused by swapping values, not rewriting code. Everything client-specific is an .env value, a content asset, or a GHL setting.', base_size=9.5, color=MUTE)
note.runs[0].italic = True
for r in note.runs:
    r.italic = True

doc.add_paragraph()

# ── Parse markdown ──────────────────────────────────────────
with open(SRC, encoding='utf-8') as f:
    lines = f.read().split('\n')

i = 0
n = len(lines)

def is_table_row(line):
    return line.strip().startswith('|')

while i < n:
    line = lines[i]
    stripped = line.strip()

    if stripped == '':
        i += 1
        continue

    # Table
    if is_table_row(stripped):
        table_lines = []
        while i < n and is_table_row(lines[i].strip()):
            table_lines.append(lines[i].strip())
            i += 1
        # remove separator row (---|---|---)
        rows = [l for l in table_lines if not re.match(r'^\|[\s:\-|]+\|$', l)]
        parsed = []
        for r in rows:
            cells = [c.strip() for c in r.strip('|').split('|')]
            parsed.append(cells)
        if parsed:
            ncols = len(parsed[0])
            t = doc.add_table(rows=len(parsed), cols=ncols)
            t.style = 'Table Grid'
            t.autofit = True
            for ridx, row in enumerate(parsed):
                for cidx, cell_text in enumerate(row):
                    if cidx >= ncols:
                        continue
                    cell = t.cell(ridx, cidx)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p = cell.paragraphs[0]
                    p.text = ''
                    add_inline_runs(p, cell_text, base_size=9.5,
                                     color=(RGBColor(0xFF, 0xFF, 0xFF) if ridx == 0 else SLATE),
                                     base_bold=(ridx == 0))
                    if ridx == 0:
                        set_cell_shading(cell, '0E1B2A')
                    elif ridx % 2 == 0:
                        set_cell_shading(cell, 'F7F4EF')
            doc.add_paragraph().space_after = Pt(4)
        continue

    # Horizontal rule
    if stripped == '---':
        i += 1
        continue

    # Headings
    m = re.match(r'^(#{1,4})\s+(.*)$', stripped)
    if m:
        level = len(m.group(1))
        text = m.group(2).strip()
        # Pull a trailing "*(their M0)*"-style annotation off into its own muted run.
        annot_m = re.match(r'^(.*?)\s*\*\(([^)]*)\)\*\s*$', text)
        annotation = annot_m.group(2) if annot_m else None
        text = (annot_m.group(1) if annot_m else text).replace('*', '')
        if level == 1:
            # The doc's single top-level title — already covered by the custom title page.
            i += 1
            continue
        elif level == 2:
            doc.add_page_break()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(19)
            run.font.color.rgb = NAVY
            run.bold = True
            if annotation:
                run2 = p.add_run('   ' + annotation)
                run2.font.size = Pt(11)
                run2.font.color.rgb = MUTE
                run2.italic = True
            add_bottom_border(p, sz=12)
            p.space_before = Pt(4)
            p.space_after = Pt(10)
        elif level == 3:
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(13.5)
            run.font.color.rgb = NAVY
            run.bold = True
            p.space_before = Pt(14)
            p.space_after = Pt(6)
        else:
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(11.5)
            run.font.color.rgb = GOLD
            run.bold = True
            p.space_before = Pt(10)
            p.space_after = Pt(4)
        i += 1
        continue

    # Checkbox list item
    m = re.match(r'^-\s+\[ \]\s+(.*)$', stripped)
    if m:
        text = m.group(1)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(3)
        box = p.add_run('☐  ')
        box.font.size = Pt(11)
        box.font.color.rgb = GOLD
        box.bold = True
        add_inline_runs(p, text, base_size=10.2, color=SLATE)
        i += 1
        continue

    # Blockquote
    if stripped.startswith('>'):
        text = stripped.lstrip('>').strip()
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(6)
        add_bottom_border(p, sz=6)
        pPr = p._p.get_or_add_pPr()
        pBdr = pPr.find(qn('w:pBdr'))
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '18')
        left.set(qn('w:space'), '8')
        left.set(qn('w:color'), LINE_GOLD)
        pBdr.append(left)
        add_inline_runs(p, text, base_size=10, color=MUTE)
        for r in p.runs:
            r.italic = True
        i += 1
        continue

    # Bare bullet (non-checkbox '-')
    m = re.match(r'^-\s+(.*)$', stripped)
    if m:
        text = m.group(1)
        p = doc.add_paragraph(style='List Bullet')
        add_inline_runs(p, text, base_size=10.2, color=SLATE)
        i += 1
        continue

    # Plain paragraph (handle *italic-wrapped* whole line)
    text = stripped
    p = doc.add_paragraph()
    if text.startswith('*') and text.endswith('*') and not text.startswith('**'):
        add_inline_runs(p, text.strip('*'), base_size=9.5, color=MUTE)
        for r in p.runs:
            r.italic = True
    else:
        add_inline_runs(p, text, base_size=10.2, color=SLATE)
    p.paragraph_format.space_after = Pt(6)
    i += 1

doc.save(OUT)
print('Saved:', OUT)
